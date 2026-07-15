import docx as docx_lib

from ufl.ingest.docx import extract


def test_docx_extract_reads_nonempty_paragraphs(tmp_path):
    source = docx_lib.Document()
    source.add_paragraph("Birinchi paragraf.")
    source.add_paragraph("")  # bo'sh paragraf o'tkazib yuborilishi kerak
    source.add_paragraph("Ikkinchi paragraf.")
    path = tmp_path / "sample.docx"
    source.save(str(path))

    document = extract(path)

    assert [b.text for b in document.blocks] == ["Birinchi paragraf.", "Ikkinchi paragraf."]
