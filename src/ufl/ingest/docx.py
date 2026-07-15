"""DOCX (Word) fayllarni ingest qilish (python-docx)."""

from __future__ import annotations

from pathlib import Path

import docx

from ufl.ingest.base import Block, Document


def extract(path: Path) -> Document:
    source = docx.Document(str(path))
    blocks = []
    for i, paragraph in enumerate(source.paragraphs):
        text = paragraph.text.strip()
        if text:
            blocks.append(Block(text=text, page=i))
    return Document(blocks=blocks)
